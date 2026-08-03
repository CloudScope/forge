from __future__ import annotations

import re
from pathlib import Path
from typing import Any


from .core.paths import ensure_runtime_dirs, paths as forge_paths

UPLOADS = forge_paths().uploads
ensure_runtime_dirs()

ALLOWED_SUFFIXES = {".md", ".txt", ".markdown", ".rst", ".pdf", ".docx", ".doc"}


def save_upload(filename: str, data: bytes) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9._-]+", "_", Path(filename).name)[:120]
    if not safe:
        safe = "requirement.md"
    path = UPLOADS / safe
    # Avoid overwrite collisions
    if path.exists():
        stem, suf = path.stem, path.suffix
        n = 1
        while path.exists():
            path = UPLOADS / f"{stem}_{n}{suf}"
            n += 1
    path.write_bytes(data)
    return path


def _extract_pdf_text(raw: bytes) -> str:
    """Extract plain text from a PDF upload (pypdf)."""
    try:
        from io import BytesIO

        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError(
            "PDF support requires pypdf. Install with: pip install pypdf"
        ) from exc

    try:
        reader = PdfReader(BytesIO(raw))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read PDF: {exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            t = ""
        if t.strip():
            parts.append(t)
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError(
            "PDF has no extractable text (scanned/image-only PDFs are not supported)"
        )
    return text


def _extract_docx_text(raw: bytes) -> str:
    """Extract plain text from a Word .docx upload (python-docx)."""
    try:
        from io import BytesIO

        from docx import Document
    except ImportError as exc:
        raise ValueError(
            "Word (.docx) support requires python-docx. Install with: pip install python-docx"
        ) from exc

    try:
        doc = Document(BytesIO(raw))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read Word document: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Word document has no extractable text")
    return text


def extract_text(path: Path, data: bytes | None = None) -> str:
    suffix = path.suffix.lower()
    raw = data if data is not None else path.read_bytes()
    if suffix == ".pdf":
        text = _extract_pdf_text(raw)
    elif suffix == ".docx":
        text = _extract_docx_text(raw)
    elif suffix == ".doc":
        raise ValueError(
            "Legacy .doc is not supported. Please save as .docx (or PDF/Markdown) and upload again."
        )
    else:
        if suffix and suffix not in ALLOWED_SUFFIXES:
            # Best-effort decode for unknown text-like uploads
            pass
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="replace")
    text = text.replace("\r\n", "\n").strip()
    if len(text) < 40:
        raise ValueError("Requirement document is too short or empty")
    if len(text) > 400_000:
        text = text[:400_000]
    return text


def summarize_document(text: str) -> dict[str, Any]:
    """Heuristic parse of a PRD into structured signals for agents."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    title = "Untitled Product"
    for ln in lines[:30]:
        if ln.startswith("#"):
            title = re.sub(r"^#+\s*", "", ln).strip()
            break
        if ln and not ln.startswith("-") and len(ln) < 120:
            title = ln
            break

    # Product short name guess
    name_match = re.search(
        r"(?:working name|product|platform)\s*[:\-]\s*([A-Za-z0-9 _-]{2,40})",
        text,
        re.I,
    )
    if name_match:
        product_name = name_match.group(1).strip().split()[0]
    else:
        # First capitalized token in title
        parts = re.findall(r"[A-Za-z][A-Za-z0-9]+", title)
        product_name = parts[0] if parts else "Product"

    # Feature keywords
    keywords = {
        "short_url": r"short\s*url|shorten|short code|tiny\s*url",
        "redirect": r"redirect|302|301",
        "custom_alias": r"alias|vanity",
        "analytics": r"analytics|click\s*event|aggregat",
        "expiration": r"expir",
        "qr_code": r"\bqr\b",
        "rate_limiting": r"rate\s*limit",
        "bulk": r"bulk",
        "preview": r"preview|ssrf",
        "admin": r"admin|api\s*key",
        "health": r"healthz|readiness|/metrics|prometheus",
        "caching": r"cache|redis",
        "auth": r"auth|api\s*key|jwt|oauth",
    }
    features: list[str] = []
    for key, pat in keywords.items():
        if re.search(pat, text, re.I):
            features.append(key)

    # FR / NFR lines
    frs = []
    nfrs = []
    for ln in lines:
        if re.match(r"^(###?\s*)?FR[-_]?\d+", ln, re.I) or re.search(
            r"\bFR[-_]?\d+\b", ln
        ):
            frs.append(re.sub(r"^#+\s*", "", ln))
        elif re.match(r"^(###?\s*)?NFR[-_]?\d+", ln, re.I) or re.search(
            r"\bNFR[-_]?\d+\b", ln
        ):
            nfrs.append(re.sub(r"^#+\s*", "", ln))
        elif re.match(r"^[-*]\s+.+(shall|must|should)\b", ln, re.I):
            frs.append(ln.lstrip("-* ").strip())

    # Sections by markdown headings
    sections: dict[str, str] = {}
    current = "preamble"
    buf: list[str] = []
    for ln in text.splitlines():
        if re.match(r"^#{1,3}\s+", ln):
            if buf:
                sections[current] = "\n".join(buf).strip()
            current = re.sub(r"^#+\s*", "", ln).strip().lower()
            buf = []
        else:
            buf.append(ln)
    if buf:
        sections[current] = "\n".join(buf).strip()

    goals = []
    for key, body in sections.items():
        if "goal" in key or "overview" in key:
            for ln in body.splitlines():
                s = ln.strip(" -*")
                if s and len(s) > 12:
                    goals.append(s)
    goals = goals[:8]

    latency = None
    m = re.search(r"p99\s*[<≤]\s*(\d+)\s*ms", text, re.I)
    if m:
        latency = int(m.group(1))
    avail = None
    m = re.search(r"(99\.9+%|\d(?:\.\d+)?\s*nines)", text, re.I)
    if m:
        avail = m.group(1)

    ambiguous = bool(
        re.search(r"\bTBD\b|\bTODO\b|to be decided|unclear|ambiguous", text, re.I)
    )

    return {
        "title": title,
        "product_name": product_name,
        "features": features,
        "fr_lines": frs[:40],
        "nfr_lines": nfrs[:20],
        "goals": goals,
        "sections": {k: v[:2000] for k, v in list(sections.items())[:20]},
        "redirect_p99_ms": latency or 50,
        "availability": avail or "99.99%",
        "ambiguous": ambiguous,
        "char_count": len(text),
        "line_count": len(lines),
        "excerpt": text[:1200],
    }
